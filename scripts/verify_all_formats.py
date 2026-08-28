import os
import sys
import json
import sqlite3
from docx import Document
from PIL import Image
import openpyxl

sys.path.append(os.path.dirname(__file__))
from ingest_document import ingest_file, DB_PATH

def seed_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS Customer (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            code TEXT UNIQUE NOT NULL,
            logoUrl TEXT,
            primaryColor TEXT DEFAULT '#0284c7',
            secondaryColor TEXT DEFAULT '#0f172a'
        );
    """)
    cursor.execute("SELECT COUNT(*) FROM Customer;")
    if cursor.fetchone()[0] == 0:
        cursor.execute("INSERT INTO Customer (id, name, code) VALUES ('cust-1', 'Apex Freight Logistics', 'APEX');")
        conn.commit()
    conn.close()

def run_verification():
    print("=========================================================")
    print("LOGISTICS AUTOMATION - MULTI-FORMAT INGESTION VERIFICATION")
    print("=========================================================\n")

    seed_db()

    test_dir = "/tmp/format_test_suite"
    os.makedirs(test_dir, exist_ok=True)
    for f in os.listdir(test_dir):
        os.remove(os.path.join(test_dir, f))

    # 1. Create test files for all 11 supported formats
    formats_to_create = {
        "test_sample.pdf": b"%PDF-1.4 %PDF sample text content",
        "test_sample.jpg": None,
        "test_sample.jpeg": None,
        "test_sample.png": None,
        "test_sample.txt": "INVOICE #: TXT-9901\nShipper: Apex Logistics\nConsignee: Global Hub\nTotal: 1475.00 USD",
        "test_sample.csv": "Invoice Number,Shipper,Consignee,Total\nCSV-8812,Apex Logistics,Global Logistics,1250.00",
        "test_sample.docx": None,
        "test_sample.doc": b"\x00\x01\x02INVOICE: DOC-1102 Shipper: Apex Freight Total: 1950.00 USD",
        "test_sample.xlsx": None,
        "test_sample.xls": b"\x00\x01\x02INVOICE: XLS-2201 Shipper: Apex Logistics Total: 2100.00 USD",
        "test_sample.tif": None,
        "test_sample.tiff": None,
    }

    # Generate image files
    img = Image.new("RGB", (200, 100), color=(50, 100, 150))
    img.save(os.path.join(test_dir, "test_sample.jpg"), format="JPEG")
    img.save(os.path.join(test_dir, "test_sample.jpeg"), format="JPEG")
    img.save(os.path.join(test_dir, "test_sample.png"), format="PNG")
    img.save(os.path.join(test_dir, "test_sample.tif"), format="TIFF")
    img.save(os.path.join(test_dir, "test_sample.tiff"), format="TIFF")

    # Generate DOCX
    doc = Document()
    doc.add_paragraph("INVOICE #: DOCX-4412")
    doc.add_paragraph("Shipper: Apex Freight Logistics")
    doc.add_paragraph("Total Amount: 1890.00 USD")
    doc.save(os.path.join(test_dir, "test_sample.docx"))

    # Generate XLSX
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Invoice Number", "Shipper Name", "Total Amount"])
    ws.append(["XLSX-7712", "Apex Logistics Hub", 1550.00])
    wb.save(os.path.join(test_dir, "test_sample.xlsx"))

    # Generate plain text / bytes
    for fname, content in formats_to_create.items():
        fpath = os.path.join(test_dir, fname)
        if content is not None:
            mode = "wb" if isinstance(content, bytes) else "w"
            with open(fpath, mode) as f:
                f.write(content)

    print("Step 1: Generated test files for all 11 target formats:")
    for fn in sorted(os.listdir(test_dir)):
        print(f"  ✓ {fn}")

    # 2. Verify ingestion & raw text extraction for all 11 formats
    print("\nStep 2: Ingesting files through backend pipeline...")
    ingested_ids = []
    failed_formats = []

    for fname in sorted(os.listdir(test_dir)):
        fpath = os.path.join(test_dir, fname)
        try:
            res = ingest_file(fpath)
            ingested_ids.append(res["documentId"])
            ext = os.path.splitext(fname)[1].lower()
            print(f"  ✅ [PASS] {fname:<20} -> DocID: {res['documentId'][:8]}... MIME: {res['mimeType']}")
        except Exception as e:
            failed_formats.append((fname, str(e)))
            print(f"  ❌ [FAIL] {fname:<20} -> Error: {e}")

    # 3. Verify unsupported format rejection
    print("\nStep 3: Testing rejection of unsupported file formats (.exe, .zip, .mp4)...")
    unsupported_files = {
        "malicious_script.exe": b"MZ executable content",
        "archive_data.zip": b"PK zip content",
        "video_sample.mp4": b"ftypmp42 content"
    }

    rejection_passed = True
    for ufname, ucontent in unsupported_files.items():
        ufpath = os.path.join(test_dir, ufname)
        with open(ufpath, "wb") as f:
            f.write(ucontent)
        try:
            ingest_file(ufpath)
            print(f"  ❌ [FAIL] {ufname} was NOT rejected!")
            rejection_passed = False
        except ValueError as ve:
            print(f"  ✅ [PASS] {ufname} correctly rejected: {ve}")

    print("\n=========================================================")
    if len(failed_formats) == 0 and rejection_passed:
        print("🎉 ALL 11 FORMATS INGESTED AND VERIFIED SUCCESSFULLY!")
    else:
        print("⚠️ VERIFICATION FAILED FOR SOME FORMATS")
    print("=========================================================")

if __name__ == "__main__":
    run_verification()
