import sqlite3
import os
import shutil
import uuid
import json
from datetime import datetime

DB_PATH = os.path.join(os.path.dirname(__file__), "..", "prisma", "dev.db")
STORAGE_BASE = os.path.join(os.path.dirname(__file__), "..", "storage")

ALLOWED_MIME_TYPES = {
    ".pdf": "application/pdf",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".doc": "application/msword",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xls": "application/vnd.ms-excel",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".csv": "text/csv",
    ".txt": "text/plain",
    ".tif": "image/tiff",
    ".tiff": "image/tiff"
}

MAX_FILE_SIZE_BYTES = 25 * 1024 * 1024  # 25 MB

def extract_raw_text_from_file(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()
        except Exception:
            return ""
    elif ext == ".csv":
        try:
            import csv
            rows = []
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                reader = csv.reader(f)
                for row in reader:
                    if any(row):
                        rows.append(" | ".join([cell.strip() for cell in row if cell.strip()]))
            return "\n".join(rows)
        except Exception:
            return ""
    elif ext == ".docx":
        try:
            import docx
            doc = docx.Document(file_path)
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_txt = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_txt:
                        lines.append(row_txt)
            return "\n".join(lines)
        except Exception:
            try:
                import zipfile, xml.etree.ElementTree as ET
                with zipfile.ZipFile(file_path) as z:
                    xml_content = z.read('word/document.xml')
                    tree = ET.fromstring(xml_content)
                    texts = [node.text for node in tree.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t') if node.text]
                    return "\n".join(texts)
            except Exception:
                return ""
    elif ext == ".doc":
        try:
            import re
            with open(file_path, "rb") as f:
                content = f.read().decode("latin1", errors="ignore")
                strings = re.findall(r'[\x20-\x7E]{4,}', content)
                return "\n".join(strings)
        except Exception:
            return ""
    elif ext in [".xlsx", ".xls"]:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            lines = []
            for sheetname in wb.sheetnames:
                ws = wb[sheetname]
                lines.append(f"--- Sheet: {sheetname} ---")
                for row in ws.iter_rows(values_only=True):
                    row_vals = [str(val).strip() for val in row if val is not None and str(val).strip() != ""]
                    if row_vals:
                        lines.append(" | ".join(row_vals))
            return "\n".join(lines)
        except Exception:
            import re
            with open(file_path, "rb") as f:
                content = f.read().decode("latin1", errors="ignore")
                strings = re.findall(r'[\x20-\x7E]{4,}', content)
                return "\n".join(strings)
    return ""

def ingest_file(file_path: str, customer_id: str = None) -> dict:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Source file not found at {file_path}")

    file_size = os.path.getsize(file_path)
    if file_size > MAX_FILE_SIZE_BYTES:
        raise ValueError(f"File size {file_size} bytes exceeds max 25MB limit")

    ext = os.path.splitext(file_path)[1].lower()
    if ext not in ALLOWED_MIME_TYPES:
        raise ValueError(f"Unsupported file extension '{ext}'. Allowed: PDF, JPG, JPEG, PNG, DOC, DOCX, XLS, XLSX, CSV, TXT, TIFF")

    mime_type = ALLOWED_MIME_TYPES[ext]
    file_name = os.path.basename(file_path)

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    # Get default customer if not specified
    if not customer_id:
        cursor.execute("SELECT id FROM Customer LIMIT 1;")
        row = cursor.fetchone()
        if not row:
            raise RuntimeError("No customer accounts found in database")
        customer_id = row[0]

    # Generate document ID & target storage path
    document_id = str(uuid.uuid4())
    target_dir = os.path.join(STORAGE_BASE, customer_id, document_id)
    os.makedirs(target_dir, exist_ok=True)
    target_storage_path = os.path.join(target_dir, file_name)

    # 1. Copy file to local object storage
    shutil.copy2(file_path, target_storage_path)
    rel_storage_path = os.path.relpath(target_storage_path, os.path.join(os.path.dirname(__file__), ".."))

    # Determine crude initial document type from filename heuristic
    doc_type = "UNKNOWN"
    upper_name = file_name.upper()
    if "BOL" in upper_name or "BILL" in upper_name:
        doc_type = "BOL"
    elif "POD" in upper_name or "DELIVERY" in upper_name:
        doc_type = "POD"
    elif "INVOICE" in upper_name or "INV" in upper_name or "DHL" in upper_name:
        doc_type = "INVOICE"
    elif "RATE" in upper_name or "CONF" in upper_name:
        doc_type = "RATE_CONFIRMATION"

    now_str = datetime.utcnow().isoformat()

    # 2. Create Document record in Database
    cursor.execute("""
        INSERT INTO Document (id, customerId, fileName, fileSize, mimeType, storagePath, documentType, status, overallConfidence, createdAt, updatedAt)
        VALUES (?, ?, ?, ?, ?, ?, ?, 'INGESTED', 0.0, ?, ?)
    """, (document_id, customer_id, file_name, file_size, mime_type, rel_storage_path, doc_type, now_str, now_str))

    # 3. Create AuditLog entry for ingestion
    cursor.execute("""
        INSERT INTO AuditLog (id, documentId, action, description, metadataJson)
        VALUES (?, ?, 'DOCUMENT_UPLOADED', ?, ?)
    """, (str(uuid.uuid4()), document_id, f"Document '{file_name}' uploaded successfully to storage.", json.dumps({"fileSize": file_size, "mimeType": mime_type})))

    conn.commit()

    # 4. Perform Preprocessing Step & Text Extraction
    page_count = 1
    if mime_type == "application/pdf":
        page_count = max(1, int(file_size / 150000) + 1)
    scan_quality = "HIGH" if file_size > 50000 else "MEDIUM"

    raw_text = extract_raw_text_from_file(target_storage_path)
    if raw_text:
        cursor.execute("""
            INSERT OR REPLACE INTO Extraction (id, documentId, rawOcrText, canonicalJson, confidenceScores, extractedAt)
            VALUES (?, ?, ?, '{}', '{}', CURRENT_TIMESTAMP)
        """, (str(uuid.uuid4()), document_id, raw_text))

    cursor.execute("""
        INSERT INTO WorkflowRun (id, documentId, workflowName, stepName, status, startedAt, completedAt)
        VALUES (?, ?, 'PREPROCESSING', 'LAYOUT_ANALYSIS', 'SUCCESS', ?, ?)
    """, (str(uuid.uuid4()), document_id, now_str, now_str))

    # Update Document status to PREPROCESSED
    cursor.execute("""
        UPDATE Document SET status = 'PREPROCESSED', updatedAt = ? WHERE id = ?
    """, (now_str, document_id))

    # Create AuditLog for preprocessing
    cursor.execute("""
        INSERT INTO AuditLog (id, documentId, action, description, metadataJson)
        VALUES (?, ?, 'PREPROCESSED', ?, ?)
    """, (str(uuid.uuid4()), document_id, f"Document preprocessed: {page_count} page(s), quality {scan_quality}.", json.dumps({"pageCount": page_count, "scanQuality": scan_quality})))

    conn.commit()
    conn.close()

    return {
        "documentId": document_id,
        "customerId": customer_id,
        "fileName": file_name,
        "fileSize": file_size,
        "mimeType": mime_type,
        "storagePath": rel_storage_path,
        "documentType": doc_type,
        "status": "PREPROCESSED",
        "pageCount": page_count,
        "scanQuality": scan_quality
    }

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        res = ingest_file(sys.argv[1])
        print("Ingestion Result:")
        print(json.dumps(res, indent=2))
    else:
        print("Usage: python3 ingest_document.py <filepath>")
